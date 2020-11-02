# -*- coding: utf-8 -*-
import sys
import os
import csv
import xlrd,xlwt
import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor
from PIL import Image,ImageDraw
from io import BytesIO


# 写入数据到CSV
def Write2CSV(file_path, head_row, content_rows):
    f = open(file_path,'w',encoding='utf-8', newline='')
    csv_writer = csv.writer(f)
    # 构建列表头
    csv_writer.writerow(head_row)
    # 写入csv文件内容
    for i in range(len(content_rows)):
        csv_writer.writerow(content_rows[i])
    f.close()

# 读取CSV文件
def ReadCSV(file_path):
    csvDatas = []
    tempCsvDatas = []
    with open(file_path,newline='',encoding='UTF-8') as csvfile:
        rows = csv.reader(csvfile)
        for row in rows:
            for i in row:
                tempCsvDatas.append(i)
            print(','.join(row))
            csvDatas.append(tempCsvDatas)
    return csvDatas

# 写入数据到Excel
def Write2Excel(file_path, sheet_name, contents):
    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet(sheet_name)

    # 设置字体字样
    style = xlwt.XFStyle()
    font = xlwt.Font()
    font.name = 'Times New Roman' 
    font.bold = True # 黑体
    font.underline = True # 下划线
    font.italic = True # 斜体字
    style.font = font # 设定样式

    # 设置单元格宽度 256*N col(0)是第一列
    # ws.col(0).width = 256*20
    # 设置单元格高度 row(0)是第一行
    # tall_style = xlwt.easyxf('font:height 720;') # 36pt,类型小初的字号 
    # ws.row(0).set_style(tall_style)

    # 输入一个日期到单元格
    # style.num_format_str = 'M/D/YY' # Other options: D-MMM-YY, D-MMM, MMM-YY, h:mm, h:mm:ss, h:mm, h:mm:ss, M/D/YY h:mm, mm:ss, [h]:mm:ss, mm:ss.0
    # ws.write(0, 4, datetime.datetime.now(), style)

    # 添加一个公式
    # ws.write(0, 4, 5) # Outputs 5
    # ws.write(0, 5, 2) # Outputs 2
    # ws.write(1, 4, xlwt.Formula('E1*F1')) # Should output "10" (A1[5] * A2[2])
    # ws.write(1, 5, xlwt.Formula('SUM(E1,F1)')) # Should output "7" (A1[5] + A2[2])

    # 添加一个超链接
    # ws.write(0, 4, xlwt.Formula('HYPERLINK("https://www.baidu.com";"BaiDu")')) # Outputs the text "BaiDu" linking to https://www.baidu.com

    # 设置对齐方式
    # alignment = xlwt.Alignment() # Create Alignment
    # alignment.horz = xlwt.Alignment.HORZ_CENTER # May be: HORZ_GENERAL, HORZ_LEFT, HORZ_CENTER, HORZ_RIGHT, HORZ_FILLED, HORZ_JUSTIFIED, HORZ_CENTER_ACROSS_SEL, HORZ_DISTRIBUTED
    # alignment.vert = xlwt.Alignment.VERT_CENTER # May be: VERT_TOP, VERT_CENTER, VERT_BOTTOM, VERT_JUSTIFIED, VERT_DISTRIBUTED
    # style = xlwt.XFStyle() # Create Style
    # style.alignment = alignment # Add Alignment to Style
    # ws.write(0, 4, 'aa', style)

    # 添加单元格边框
    # borders = xlwt.Borders() # Create Borders
    # borders.left = xlwt.Borders.DASHED # DASHED虚线 NO_LINE没有 THIN实线
        
    # # May be: NO_LINE, THIN, MEDIUM, DASHED, DOTTED, THICK, DOUBLE, HAIR, MEDIUM_DASHED, THIN_DASH_DOTTED, MEDIUM_DASH_DOTTED, THIN_DASH_DOT_DOTTED, MEDIUM_DASH_DOT_DOTTED, SLANTED_MEDIUM_DASH_DOTTED, or 0x00 through 0x0D.
    # borders.right = xlwt.Borders.DASHED
    # borders.top = xlwt.Borders.DASHED
    # borders.bottom = xlwt.Borders.DASHED
    # borders.left_colour = 0x40
    # borders.right_colour = 0x40
    # borders.top_colour = 0x40
    # borders.bottom_colour = 0x40
    # style = xlwt.XFStyle() # Create Style
    # style.borders = borders # Add Borders to Style
    # ws.write(0, 4, 'bb', style)

    # 添加背景色
    pattern = xlwt.Pattern() # Create the Pattern
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN # May be: NO_PATTERN, SOLID_PATTERN, or 0x00 through 0x12
    pattern.pattern_fore_colour = 5 # May be: 8 through 63. 0 = Black, 1 = White, 2 = Red, 3 = Green, 4 = Blue, 5 = Yellow, 6 = Magenta, 7 = Cyan, 16 = Maroon, 17 = Dark Green, 18 = Dark Blue, 19 = Dark Yellow , almost brown), 20 = Dark Magenta, 21 = Teal, 22 = Light Gray, 23 = Dark Gray, the list goes on...
    style = xlwt.XFStyle() # Create the Pattern
    style.pattern = pattern # Add Pattern to Style
    ws.write(0, 4, 'cc', style)

    for i in range(len(contents)):
        for j in range(len(contents[i])):
            ws.write(i, j, label=contents[i][j]) # 不带样式的写入
            # ws.write(i, j, contents[i][j], style) # 带样式的写入

    wb.save(file_path)

# 写入数据到Word
def Write2Word(file_path, headline):
    document = Document()
    # 段落集合 document.paragraphs
    # 表格集合 document.tables
    # 节集合 document.sections
    # 样式集合 document.styles
    # 内置图形 document.inline_shapes
    
    document.add_heading(u'MS WORD写入测试',0)
    document.add_heading(u'一级标题',1)
    document.add_heading(u'二级标题',2)

    #设置字号
    paragraph = document.add_paragraph(u'我们在做文本测试！')
    run = paragraph.add_run(u'设置字号、')
    run.font.size = Pt(24)

    #设置字体
    run = paragraph.add_run('Set Font,')
    run.font.name = 'Consolas'

    #设置中文字体
    run = paragraph.add_run(u'设置中文字体、')
    run.font.name=u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    #设置颜色
    run.font.color.rgb = RGBColor(0x0, 0xff, 0x0) # 绿色

    #设置斜体
    run = paragraph.add_run(u'斜体、')
    run.italic = True

    #设置粗体
    run = paragraph.add_run(u'粗体')
    run.bold = True

    #增加引用
    document.add_paragraph('Intense quote', style='Intense Quote')


    #增加无序列表
    document.add_paragraph(
        u'无序列表元素1', style='List Bullet'
    )
    document.add_paragraph(
        u'无序列表元素2', style='List Bullet'
    )
    #增加有序列表
    document.add_paragraph(
        u'有序列表元素1', style='List Number'
    )
    document.add_paragraph(
        u'有序列表元素2', style='List Number'
    )

    #增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
    img_name = "F:/python-libs/trunk/image/combine_image/need_combine_images/test1.jpg"
    document.add_picture(img_name, width=Inches(1.25))

    #增加表格
    table = document.add_table(rows=1, cols=3)
    # 新增表格的行,列
    # table.add_row()
    # table.add_column(20) # 必须声明width,否则报错

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    #再增加3行表格元素
    for i in range(3):
        row_cells = table.add_row().cells
        row_cells[0].text = 'test'+str(i)
        row_cells[1].text = str(i)
        row_cells[2].text = 'desc'+str(i)

    #增加分页
    document.add_page_break()

    #添加255个圆圈到word
    p = document.add_paragraph()
    r = p.add_run()
    img_size = 20
    for x in range(255):
        im = Image.new("RGB",(img_size,img_size),"white")
        draw_obj = ImageDraw.Draw(im)
        draw_obj.ellipse((0,0,img_size-1,img_size-1),fill=255-x) # 画圆
        fake_buf_file = BytesIO() # 用BytesIO将图片保存在内存中,减少磁盘操作
        im.save(fake_buf_file,"png")
        r.add_picture(fake_buf_file) # 插入图片
        fake_buf_file.close()

    # 定义标题,再添加内容
    # p_total = document.add_heading()
    # r_total = p_total.add_run("执行结果如下：")
    # r_total.font.bold = True # 字体加粗

    # # 添加表格
    # table = document.add_table(rows=1, cols=3, style="Light List Accent 5")
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'testName'
    # hdr_cells[1].text = 'param'
    # hdr_cells[2].text = 'exc'

    #添加二级标题
    # p_total = document.add_heading("", 2)
    # r_total = p_total.add_run("this is second headline")
    # r_total.font.bold = True

    # 添加图片
    # img_name = "F:/python-libs/trunk/image/combine_image/need_combine_images/test1.jpg"
    # document.add_picture(img_name, width=Inches(1.5))

    # 插入有序表
    document.add_paragraph('time') # 增加一个paragraph

    # #插入有序列表,段落的前面会有序号123
    # document.add_paragraph('把冰箱门打开',style='List Number')
    # document.add_paragraph('把大象装进去',style='List Number')
    # document.add_paragraph('把冰箱门关上',style='List Number')

    # #插入无序列表，段落的前面没有序号
    # document.add_paragraph('把冰箱门打开',style='List Bullet')
    # document.add_paragraph('把大象装进去',style='List Bullet')
    # document.add_paragraph('把冰箱门关上',style='List Bullet')

    # 插入一个6行6列的表格
    # table=document.add_table(rows=6,cols=6,style='Table Grid')
    # for i in range(0,6):
    #     for j in range(0,6):
    #         table.cell(i,j).text="第{i}行{j}列".format(i=i+1,j=j+1)



    document.save(file_path)  # 保存文档

# 读取word数据
def ReadWord(file_path):
    #打开文档
    document = Document(u''+file_path)
    #读取每段资料
    l = [ paragraph.text.encode('gb2312') for paragraph in document.paragraphs];
    #输出并观察结果，也可以通过其他手段处理文本即可
    for i in l:
        print (i)
    #读取表格材料，并输出结果
    tables = [table for table in document.tables];
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text.encode('gb2312'),'\t', print)
        print ('\n')

# 写入数据到txt
def Write2Txt(file_path, contents):
    f = open(file_path, mode='a+', encoding='utf-8')
    for i in range(len(contents)):
        f.write(contents[i] + '\n')
    f.close()


if __name__ == '__main__':
    # 切换到当前执行文件的工作目录下
    filename = sys.argv[0]
    dirname = os.path.dirname(filename)
    abspath = os.path.abspath(dirname)
    os.chdir(abspath)

    file_path = abspath + "/test.csv"
    Write2CSV(file_path, ["姓名","年龄","性别"], [["l",'18','男'],["c",'20','男'],["w",'22','女']])
    ReadCSV(file_path)

    file_path = abspath + "/test.txt"
    # Write2Txt(file_path, ["aaaa","bbbb"])

    file_path = abspath + "/test.xls"
    # Write2Excel(file_path, "test", [["a","b"],["1","2"],["4","5"]])

    file_path = abspath + "/test.docx"
    # Write2Word(file_path, "this is headline")
    # ReadWord(file_path)